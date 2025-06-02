#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import pandas as pd
import numpy as np
from scipy.stats import chi2_contingency, f_oneway, ttest_ind
from pathlib import Path
from docx import Document
from docx.shared import Inches

# === STATS FUNCTIONS ===
def cohen_d(x, y):
    nx, ny = len(x), len(y)
    if nx < 2 or ny < 2:
        return np.nan
    pooled_std = np.sqrt(((nx - 1) * x.std(ddof=1) ** 2 + (ny - 1) * y.std(ddof=1) ** 2) / (nx + ny - 2))
    if pooled_std == 0:
        return np.nan
    return (x.mean() - y.mean()) / pooled_std

def cramers_v(confusion_matrix):
    chi2, _, _, _ = chi2_contingency(confusion_matrix)
    n = confusion_matrix.sum().sum()
    r, k = confusion_matrix.shape
    return np.sqrt(chi2 / (n * (min(k - 1, r - 1))))

# === CONFIGURATION ===
meta_file = Path("results/filtered/meta.tsv")
surveil_file = Path("results/filtered/surveillance.tsv")
anthro_file = Path("results/filtered/anthro.tsv")

REFEED_COL = "Feed"

meta = pd.read_csv(meta_file, sep="\t", index_col=0)
meta = meta.dropna(subset=[REFEED_COL])
surveil = pd.read_csv(surveil_file, sep="\t", index_col=0)

df = meta.join(surveil)

categorical_cols = [
        'Recovery',
        'Sex',
        'Delivery_Mode',
        'PoB']
continuous_cols = [
        "BF",
        'fail_no.failure',
        'days_of_catchup'] 

# One-hot encode categorical vars
df_encoded = pd.get_dummies(df[categorical_cols])

# Drop redundant onehot columns
df_encoded = df_encoded.drop(['Recovery_No recovery', 'Sex_Male', 'Delivery_Mode_Vaginal'], axis=1)

df_combined = pd.concat([df[[REFEED_COL] + continuous_cols], df_encoded], axis=1)

groups = df_combined[REFEED_COL].unique()
grouped = df_combined.groupby(REFEED_COL)

summary_rows = []

# === CATEGORICAL (ONE-HOT) VARIABLES ===
for col in df_encoded.columns:
    row = {"Variable": col}

    for group in groups:
        vals = df_combined[df_combined[REFEED_COL] == group][col]
        mean = vals.mean() * 100
        row[group] = f"{mean:.1f}% (n={len(vals)})"

    try:
        # Create contingency table
        contingency = []
        for group in groups:
            counts = df_combined[df_combined[REFEED_COL] == group][col].value_counts()
            # Ensure consistent indexing for False (0) and True (1)
            false_count = counts.get(0, counts.get(False, 0))
            true_count = counts.get(1, counts.get(True, 0))
            contingency.append([false_count, true_count])

        table = pd.DataFrame(contingency, index=groups, columns=["False", "True"])
        print(f"Contingency table for {col}:\n{table}")  # Debug output

        # Check for valid contingency table
        if table.sum().sum() == 0 or table.eq(0).all().any():
            row["P-value"] = "NA"
            row["Effect size"] = "NA"
            print(f"Warning: Invalid contingency table for {col}")
            continue

        # Perform chi-square test
        chi2, p, _, expected = chi2_contingency(table)

        # Check expected frequencies
        if (expected < 5).any():
            row["P-value"] = "NA"
            row["Effect size"] = "NA"
            print(f"Warning: Low expected frequencies (<5) for {col}. Expected:\n{expected}")
            continue

        # Calculate Cramér’s V
        v = cramers_v(table)
        row["P-value"] = f"{p:.4f}"
        row["Effect size"] = f"{v:.2f}"

    except Exception as e:
        row["P-value"] = "NA"
        row["Effect size"] = "NA"
        print(f"Error processing {col}: {str(e)}")

    summary_rows.append(row)

# === CONTINUOUS VARIABLES ===
for col in continuous_cols:
    if col not in df_combined.columns:
        continue
    row = {"Variable": col}
    group_vals = []

    for group in groups:
        vals = df_combined[df_combined[REFEED_COL] == group][col].dropna()
        group_vals.append(vals)
        row[group] = f"{vals.mean():.2f} ± {vals.std():.2f} (n={len(vals)})"

    try:
        if len(groups) == 2:
            stat, p = ttest_ind(*group_vals, equal_var=False)
            d = cohen_d(*group_vals)
            row["Effect size"] = f"{d:.2f}"
        else:
            stat, p = f_oneway(*group_vals)
            row["Effect size"] = "NA"
        row["P-value"] = f"{p:.4f}"
    except Exception:
        row["P-value"] = "NA"
        row["Effect size"] = "NA"

    summary_rows.append(row)

# === FORMAT TABLE ===
summary_df = pd.DataFrame(summary_rows)
summary_df.set_index("Variable", inplace=True)

# Rename columns with (n=...) format
column_renames = {
    group: f"{group} (n={len(df_combined[df_combined[REFEED_COL] == group])})"
    for group in groups
}
summary_df.rename(columns=column_renames, inplace=True)

# === SAVE TSV ===
tsv_output = Path("results/refeed_summary_with_stats.tsv")
summary_df.to_csv(tsv_output, sep="\t")
print(f"Saved TSV summary to {tsv_output}")

# === EXPORT TO DOCX ===
doc = Document()
doc.add_heading("Summary of Participant Characteristics by Refeeding Method", level=1)

doc.add_paragraph(
    "Effect size: Cohen’s d for continuous variables, Cramér’s V for binary categorical variables. "
    "Values are shown as mean ± SD (n) or % (n)."
)

table = doc.add_table(rows=1, cols=len(summary_df.columns) + 1)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Variable'
for i, col in enumerate(summary_df.columns):
    hdr_cells[i + 1].text = col

# Data rows
for idx, row in summary_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(idx)
    for i, val in enumerate(row):
        row_cells[i + 1].text = str(val)

# Save Word document
word_output = Path("results/refeed_summary_with_stats.docx")
doc.save(word_output)
print(f"Saved DOCX summary to {word_output}")

