#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Author: Theo Portlock
"""

import numpy as np
import pandas as pd
from docx import Document

# For samplesheet processing
docs = []
docs.append(Document('data/InoliWD_Batch3_JOS_CS-ANZ DNA Sample Information Form.docx'))
docs.append(Document('data/InoliWD_Batch4_JOS_CS-ANZ DNA Sample Information Form.docx'))
docs.append(Document('data/InoliWD_Batch5_JOS_CS-ANZ DNA Sample Information Form.docx'))
docs.append(Document('data/InoliWD_JOS_CS-ANZ DNA Sample Information Form.docx'))
docs.append(Document('data/CS-ANZ DNA Sample Information Form_Hui Hui_20240220_TPupdated.docx'))

output = []

for i, doc in enumerate(docs):
    # Extract the first table from the document
    table = doc.tables[0]
    coltext = []
    # Process rows starting from row 12 onwards
    for row in table.rows[12:]:
        rowtext = []
        for cell in row.cells:
            rowtext.append(cell.text)
        coltext.append(rowtext)

    df = pd.DataFrame(coltext)
    df = df.T.drop_duplicates().T
    df = df.iloc[:, :-1]  # remove remark column
    df.columns = df.iloc[0]
    df = df.iloc[1:, :]
    df = df.set_index('*Name on tube')
    df = df.drop_duplicates()
    output.append(df)

# Combine dataframes and perform cleaning
df = pd.concat(output).sort_index().drop_duplicates().iloc[1:]
df = (
    df['Name on report']
    .drop_duplicates()
    .to_frame()
    .reset_index()
    .set_index('Name on report')
    .drop_duplicates()
)  # remove duplicate, can be used for future batch effect analysis

# Remove controls by filtering index
df = df.loc[df.index.str.startswith('LCC')]

# Assign batch numbers based on the starting characters of '*Name on tube'
df.loc[(df['*Name on tube'].str[:2] == 'C1') | (df['*Name on tube'].str[:2] == 'C2'), 'batch'] = 1
df.loc[df['*Name on tube'].str[:2] == 'C5', 'batch'] = 2
df.loc[df['*Name on tube'].str[:1] == 'J', 'batch'] = 3
df.loc[df['*Name on tube'].str[:1] == 'M', 'batch'] = 4
df.loc[df['*Name on tube'].str[:1] == 'S', 'batch'] = 5

# Create an ID and assign timepoints based on patterns in the index
df['ID'] = df.index.str[:7]
df.loc[df.index.str[7:] == '1001', 'timepoint'] = '000'
df.loc[df.index.str[3] == '3', 'timepoint'] = '104'
df.loc[df.index.str[7:] == '1002', 'timepoint'] = '012'
df.loc[df.index.str.contains(r'2\d\d\d1002'), 'timepoint'] = '052'
df.loc[df.index.str[7:] == '1003', 'timepoint'] = '052'

# Finalize the samplesheet
samplesheet = df.dropna()
samplesheet = samplesheet.rename(columns={'*Name on tube': 'Seq_ID'})
samplesheet.index = samplesheet.ID + samplesheet.timepoint

# Rename
samplesheet = samplesheet.rename(columns={'ID':'subjectID'})
samplesheet.timepoint = samplesheet.timepoint.astype(int)
samplesheet.to_csv('results/samplesheet.tsv', sep='\t', index=False)
